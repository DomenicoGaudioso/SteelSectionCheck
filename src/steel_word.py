from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


def _set_style(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.color.rgb = RGBColor(25, 25, 25)


def _add_title(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    document.add_paragraph(f"Documento generato il {datetime.now():%d/%m/%Y %H:%M}.")


def _add_table(document: Document, title: str, dataframe) -> None:
    document.add_heading(title, level=2)
    table = document.add_table(rows=1, cols=len(dataframe.columns))
    table.style = "Table Grid"

    header = table.rows[0].cells
    for idx, column in enumerate(dataframe.columns):
        header[idx].text = str(column)
        for run in header[idx].paragraphs[0].runs:
            run.bold = True

    for _, row in dataframe.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            if isinstance(value, float):
                cells[idx].text = f"{value:.3g}"
            else:
                cells[idx].text = str(value)


def _add_figure(document: Document, figure) -> None:
    if figure is None:
        return
    image = BytesIO()
    figure.savefig(image, format="png", dpi=170, bbox_inches="tight")
    image.seek(0)
    document.add_heading("1.4 Schema grafico", level=2)
    document.add_picture(image, width=Inches(5.8))


def genera_word_steel(
    titolo: str,
    modello: str,
    input_df=None,
    properties_df=None,
    checks_df=None,
    figure=None,
    note: str | None = None,
) -> bytes:
    document = Document()
    _set_style(document)
    _add_title(document, titolo)

    document.add_heading("1. Relazione tecnica di calcolo", level=1)
    document.add_paragraph(
        "La relazione riporta le proprieta' geometriche della sezione in acciaio e, "
        "quando disponibili, le verifiche elastiche di tensione. Le grandezze sono "
        "espresse nel sistema mm, kN, kNm e MPa."
    )
    document.add_paragraph(f"Modello: {modello}.")
    if note:
        document.add_paragraph(note)

    if input_df is not None and not input_df.empty:
        _add_table(document, "1.1 Dati di input", input_df)
    if properties_df is not None and not properties_df.empty:
        _add_table(document, "1.2 Proprieta' geometriche", properties_df)
    if checks_df is not None and not checks_df.empty:
        _add_table(document, "1.3 Verifiche", checks_df)
    _add_figure(document, figure)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
